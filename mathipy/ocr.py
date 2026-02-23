"""Multimodal OCR for extracting text and math from assessment content."""

from __future__ import annotations

import io
import json
import logging
import re
from pathlib import Path
from typing import Any

from mathipy._api import VisionAPIClient, image_extensions, max_file_size
from mathipy.utils import extract_math_expressions, extract_numbers, extract_variables

logger = logging.getLogger(__name__)

try:
    import httpx  # noqa: F401
    httpx_available = True
except ImportError:
    httpx_available = False
    logger.warning("httpx not available - install with: pip install httpx")

try:
    from docx import Document as DocxDocument
    docx_available = True
except ImportError:
    docx_available = False

try:
    import pdfplumber
    pdfplumber_available = True
except ImportError:
    pdfplumber_available = False

system_prompt = """You are an expert OCR system for multi-modal content extraction.
Extract ALL visible text from the image accurately, including:
- Question text and instructions
- Mathematical expressions and equations (use LaTeX notation)
- Answer choices (A, B, C, D, etc.)
- Diagram labels and annotations
- Any numbers, variables, or symbols
- Table data, chart labels, and data values
- Headers, titles, and captions

IMPORTANT: You MUST always provide meaningful output.
- If the image contains text: extract all text accurately.
- If the image contains ONLY a picture with NO text: set "content_type" to "image_only"
  and provide a detailed description of what the image shows in the "image_description" field.
  Describe the image using K-12 math terms (shapes, colors, measurements, spatial relationships).
- If the image contains BOTH text and pictures: set "content_type" to "mixed",
  extract all text AND describe the visual elements.

NEVER return empty results. Every image has content worth describing.

Respond with a JSON object containing:
{
    "content_type": "text_only|image_only|mixed",
    "full_text": "complete extracted text (empty string if image_only)",
    "image_description": "detailed description of visual content using K-12 math vocabulary (shapes, colors, width, height, spatial relationships). Always populated for image_only and mixed content.",
    "question_text": "main question/problem statement (if applicable)",
    "math_expressions": ["list of mathematical expressions in LaTeX"],
    "answer_choices": {"A": "choice A text", "B": "choice B text", ...},
    "data_elements": ["table rows, chart values, or other structured data"],
    "labels": ["list of diagram/figure labels"],
    "numbers_found": ["list of numbers"],
    "variables_found": ["list of variables"],
    "extraction_confidence": 0.0-1.0
}"""

user_prompt = """Extract all text and mathematical content from this assessment image.
Be thorough and accurate. Include all visible text, equations, and symbols.
If the image contains no extractable text (picture only), describe the visual content
in detail using K-12 math vocabulary (shapes, colors, dimensions, spatial relationships).
You must ALWAYS return meaningful content - never empty results."""

describe_system_prompt = """You are an expert image description system for K-12 math education.
Provide a concise description of the image content, focusing on:
- Mathematical elements (equations, expressions, symbols)
- Geometric shapes and spatial relationships
- Diagrams, graphs, charts, and their key features
- Colors, labels, and annotations

Respond with a JSON object containing:
{
    "content_type": "image_only|mixed|text_only",
    "full_text": "",
    "image_description": "concise description here",
    "question_text": "",
    "math_expressions": [],
    "answer_choices": {},
    "data_elements": [],
    "labels": [],
    "numbers_found": [],
    "variables_found": [],
    "extraction_confidence": 0.0-1.0
}"""

describe_user_prompt_template = """Describe this image in {max_words} words or fewer.
Focus on mathematical content, shapes, diagrams, and spatial relationships.
Be concise but precise. Return the description in the JSON format specified."""


def _parse_response(response_text: str) -> dict[str, Any]:
    cleaned = re.sub(r"```(?:json)?\s*", "", response_text).strip()
    json_match = re.search(r"\{[\s\S]*\}", cleaned)
    if json_match:
        try:
            parsed = json.loads(json_match.group())
            if not parsed.get("content_type"):
                has_text = bool(parsed.get("full_text", "").strip())
                has_desc = bool(parsed.get("image_description", "").strip())
                if has_text and has_desc:
                    parsed["content_type"] = "mixed"
                elif has_text:
                    parsed["content_type"] = "text_only"
                else:
                    parsed["content_type"] = "image_only"
            return parsed
        except json.JSONDecodeError:
            pass

    logger.warning("Could not parse JSON from API response; falling back to plain-text extraction")
    text = response_text.strip()
    has_extractable_text = bool(re.search(r"[a-zA-Z0-9]{3,}", text))

    return {
        "content_type": "text_only" if has_extractable_text else "image_only",
        "full_text": text if has_extractable_text else "",
        "image_description": text if not has_extractable_text else "",
        "question_text": text if has_extractable_text else "",
        "math_expressions": extract_math_expressions(text),
        "answer_choices": _extract_answer_choices(text),
        "data_elements": [],
        "labels": [],
        "numbers_found": extract_numbers(text),
        "variables_found": extract_variables(text),
        "extraction_confidence": 0.5,
    }



def _extract_answer_choices(text: str) -> dict[str, str]:
    choices = {}
    choice_pattern = r"([A-E])[.)\s]+(.+?)(?=[A-E][.)\s]|$)"
    matches = re.findall(choice_pattern, text, re.MULTILINE | re.DOTALL)
    for letter, content in matches:
        choices[letter] = content.strip()
    return choices


def _empty_result() -> dict[str, Any]:
    return {
        "content_type": "text_only",
        "full_text": "",
        "image_description": "",
        "question_text": "",
        "math_expressions": [],
        "answer_choices": {},
        "data_elements": [],
        "labels": [],
        "numbers_found": [],
        "variables_found": [],
        "extraction_confidence": 0.0,
    }


class MultimodalOCR(VisionAPIClient):
    """Extract text and math expressions from images using vision LLMs.

    Supports Gemini and OpenAI providers. Accepts image files, URLs, bytes,
    PDFs, DOCX, and text files.

    Requires ``pip install mathipy[ocr]`` and a ``GEMINI_API_KEY`` or
    ``OPENAI_API_KEY`` in your ``.env`` file.
    """

    def extract(
        self,
        source: str | Path | bytes,
        mode: str = "full",
        max_words: int = 30,
    ) -> dict[str, Any]:
        """Extract text and math content from the given source.

        Args:
            source: Image path, URL, bytes, PDF, DOCX, or text file.
            mode: ``"full"`` for complete OCR extraction or ``"describe"`` for brief image description.
            max_words: Maximum words for describe mode (default 30).

        Returns:
            Dictionary with ``content_type``, ``full_text``, ``image_description``,
            ``question_text``, ``math_expressions``, ``answer_choices``, and
            ``extraction_confidence``.
        """
        if mode not in ("full", "describe"):
            raise ValueError(f"Unsupported mode: {mode}. Use 'full' or 'describe'.")

        if isinstance(source, bytes):
            return self._extract_from_image(source, mode, max_words)

        source_str = str(source)

        if source_str.startswith(("http://", "https://")):
            return self._extract_from_image(source_str, mode, max_words)

        path = Path(source_str)
        suffix = path.suffix.lower()

        if suffix == ".txt":
            return self._extract_from_txt(path)

        if suffix == ".docx":
            return self._extract_from_docx(path, mode, max_words)

        if suffix == ".pdf":
            return self._extract_from_pdf(path, mode, max_words)

        if suffix in image_extensions:
            return self._extract_from_image(path, mode, max_words)

        return self._extract_from_image(path, mode, max_words)

    def _extract_from_image(
        self,
        source: str | Path | bytes,
        mode: str = "full",
        max_words: int = 30,
    ) -> dict[str, Any]:
        image_b64, mime_type = self._prepare_image(source)

        if mode == "describe":
            sys_prompt = describe_system_prompt
            usr_prompt = describe_user_prompt_template.format(max_words=max_words)
        else:
            sys_prompt = system_prompt
            usr_prompt = user_prompt

        if self.provider == "gemini":
            response_text = self._call_gemini(
                image_b64, mime_type,
                system_prompt=sys_prompt,
                user_prompt=usr_prompt,
            )
        else:
            response_text = self._call_openai(
                image_b64, mime_type,
                system_prompt=sys_prompt,
                user_prompt=usr_prompt,
            )

        return _parse_response(response_text)

    def _extract_from_txt(self, path: Path) -> dict[str, Any]:
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        text = path.read_text(encoding="utf-8")
        result = _empty_result()
        result["full_text"] = text
        result["question_text"] = text.strip().split("\n")[0] if text.strip() else ""
        result["math_expressions"] = extract_math_expressions(text)
        result["answer_choices"] = _extract_answer_choices(text)
        result["numbers_found"] = extract_numbers(text)
        result["variables_found"] = extract_variables(text)
        result["extraction_confidence"] = 1.0
        result["content_type"] = "text_only"
        return result

    def _extract_from_docx(
        self, path: Path, mode: str, max_words: int
    ) -> dict[str, Any]:
        if not docx_available:
            raise ImportError(
                "python-docx is required for .docx extraction. "
                "Install with: pip install mathipy[documents]"
            )
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        if path.stat().st_size > max_file_size:
            raise ValueError(
                f"File too large ({path.stat().st_size / 1024 / 1024:.1f} MB). "
                f"Maximum allowed: {max_file_size / 1024 / 1024:.0f} MB."
            )

        doc = DocxDocument(str(path))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))

        text = "\n".join(paragraphs)
        if table_rows:
            text += "\n" + "\n".join(table_rows)

        result = _empty_result()
        result["full_text"] = text
        result["question_text"] = paragraphs[0] if paragraphs else ""
        result["math_expressions"] = extract_math_expressions(text)
        result["answer_choices"] = _extract_answer_choices(text)
        result["numbers_found"] = extract_numbers(text)
        result["variables_found"] = extract_variables(text)
        result["data_elements"] = table_rows
        result["extraction_confidence"] = 0.95
        result["content_type"] = "text_only"

        images = self._extract_docx_images(doc)
        if images:
            image_results = self._process_embedded_images(images, mode, max_words)
            if image_results:
                result = self._merge_image_results(result, image_results)

        return result

    def _extract_from_pdf(
        self, path: Path, mode: str, max_words: int
    ) -> dict[str, Any]:
        if not pdfplumber_available:
            raise ImportError(
                "pdfplumber is required for .pdf extraction. "
                "Install with: pip install mathipy[documents]"
            )
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        if path.stat().st_size > max_file_size:
            raise ValueError(
                f"File too large ({path.stat().st_size / 1024 / 1024:.1f} MB). "
                f"Maximum allowed: {max_file_size / 1024 / 1024:.0f} MB."
            )

        all_text = []
        all_images: list[tuple[bytes, str]] = []

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    all_text.append(page_text)

                for img in page.images:
                    try:
                        x0 = img["x0"]
                        top = img["top"]
                        x1 = img["x1"]
                        bottom = img["bottom"]
                        cropped = page.crop((x0, top, x1, bottom))
                        pil_img = cropped.to_image(resolution=150).original
                        buf = io.BytesIO()
                        pil_img.save(buf, format="PNG")
                        all_images.append((buf.getvalue(), "image/png"))
                    except Exception as e:
                        logger.debug(f"Could not extract PDF image: {e}")

        text = "\n".join(all_text)
        result = _empty_result()
        result["full_text"] = text
        result["question_text"] = text.strip().split("\n")[0] if text.strip() else ""
        result["math_expressions"] = extract_math_expressions(text)
        result["answer_choices"] = _extract_answer_choices(text)
        result["numbers_found"] = extract_numbers(text)
        result["variables_found"] = extract_variables(text)
        result["extraction_confidence"] = 0.95
        result["content_type"] = "text_only"

        if all_images:
            image_results = self._process_embedded_images(all_images, mode, max_words)
            if image_results:
                result = self._merge_image_results(result, image_results)

        return result

    def _extract_docx_images(
        self, doc: DocxDocument
    ) -> list[tuple[bytes, str]]:
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    images.append((image_part.blob, image_part.content_type))
                except Exception as e:
                    logger.debug(f"Could not extract docx image: {e}")
        return images

    def _process_embedded_images(
        self,
        images: list[tuple[bytes, str]],
        mode: str,
        max_words: int,
    ) -> list[dict[str, Any]]:
        results = []
        for img_bytes, _mime_type in images:
            try:
                result = self._extract_from_image(img_bytes, mode, max_words)
                results.append(result)
            except Exception as e:
                logger.debug(f"Could not process embedded image: {e}")
        return results

    def _merge_image_results(
        self,
        base: dict[str, Any],
        image_results: list[dict[str, Any]],
    ) -> dict[str, Any]:
        base["content_type"] = "mixed"

        descriptions = []
        for i, img_result in enumerate(image_results, 1):
            desc = img_result.get("image_description", "")
            if desc:
                descriptions.append(f"[Image {i}] {desc}")

        existing_desc = base.get("image_description", "")
        all_descriptions = ([existing_desc] if existing_desc else []) + descriptions
        base["image_description"] = "\n".join(all_descriptions)

        math_set = set(base.get("math_expressions", []))
        labels_set = set(base.get("labels", []))
        numbers_set = set(base.get("numbers_found", []))
        variables_set = set(base.get("variables_found", []))
        data_elements = list(base.get("data_elements", []))

        for img_result in image_results:
            math_set.update(img_result.get("math_expressions", []))
            labels_set.update(img_result.get("labels", []))
            numbers_set.update(img_result.get("numbers_found", []))
            variables_set.update(img_result.get("variables_found", []))
            data_elements.extend(img_result.get("data_elements", []))
            base.get("answer_choices", {}).update(img_result.get("answer_choices", {}))

        base["math_expressions"] = list(math_set)
        base["labels"] = list(labels_set)
        base["numbers_found"] = list(numbers_set)
        base["variables_found"] = list(variables_set)
        base["data_elements"] = data_elements

        text_confidence = base.get("extraction_confidence", 0.95)
        img_confidences = [
            r.get("extraction_confidence", 0.5) for r in image_results
        ]
        avg_img_confidence = sum(img_confidences) / len(img_confidences) if img_confidences else 0.5
        base["extraction_confidence"] = round(
            0.7 * text_confidence + 0.3 * avg_img_confidence, 2
        )

        return base
