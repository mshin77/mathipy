"""Split a multi-item document into per-item records.

Released test documents place many items in one file, marked by an identifier
line. Whether that line precedes or follows the block it names varies by
publisher, and both readings parse without error, so a wrong choice is silent.
``check_alignment`` measures which reading is right before an analysis depends
on it.
"""

from __future__ import annotations

import re
from collections.abc import Sequence
import io
from pathlib import Path
from typing import Any

from mathipy._api import _optional_import

_docx_mod, docx_available = _optional_import("docx", "python-docx")
DocxDocument = getattr(_docx_mod, "Document", None)

label_positions = ("leading", "trailing")

_stopwords = set(
    "a an the of to in for is are and or on with each following shown show "
    "which what this that from at by as be its it".split()
)


def _tokens(text: str) -> set[str]:
    return {w for w in re.findall(r"[a-z]+", text.lower())
            if len(w) > 3 and w not in _stopwords}


def _overlap(text: str, reference: str) -> float:
    target = _tokens(reference)
    return len(_tokens(text) & target) / len(target) if target else 0.0


_M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def _fuse(parts) -> str:
    joined = ""
    for part in parts:
        if joined and part and joined[-1].isalnum() and part[0].isalnum():
            joined += " "
        joined += part
    return joined


def _omml_child(el, name: str) -> str:
    return "".join(_omml_text(c) for c in el if c.tag == _M + name)


def _omml_text(el) -> str:
    """Plain text for one OMML element, with the operator each structure implies."""
    tag = el.tag[len(_M):] if el.tag.startswith(_M) else el.tag
    if tag.endswith("Pr"):
        return ""
    if tag == "t":
        return el.text or ""
    if tag == "f":
        return f"{_omml_child(el, 'num')}/{_omml_child(el, 'den')}"
    if tag == "sSup":
        return f"{_omml_child(el, 'e')}^{_omml_child(el, 'sup')}"
    if tag == "sSub":
        return f"{_omml_child(el, 'e')}_{_omml_child(el, 'sub')}"
    if tag == "sSubSup":
        return (f"{_omml_child(el, 'e')}_{_omml_child(el, 'sub')}"
                f"^{_omml_child(el, 'sup')}")
    if tag == "rad":
        return f"sqrt({_omml_child(el, 'e')})"
    if tag == "d":
        return f"({_omml_child(el, 'e')})"
    if tag == "nary":
        return (f"{_omml_child(el, 'sub')}..{_omml_child(el, 'sup')} "
                f"{_omml_child(el, 'e')}")
    return _fuse(_omml_text(c) for c in el)


def paragraph_text(para) -> str:
    """Paragraph text including equation objects, which ``Paragraph.text`` omits."""
    from docx.oxml.ns import qn

    run, math, para_math, link = qn("w:r"), _M + "oMath", _M + "oMathPara", qn("w:hyperlink")
    text_tag = qn("w:t")
    parts = []
    for child in para._p:
        if child.tag in (run, link):
            text, spaced = "".join(t.text or "" for t in child.iter(text_tag)), False
        elif child.tag in (math, para_math):
            text, spaced = _omml_text(child), True
        else:
            continue
        if not text:
            continue
        gap = (spaced or (parts and parts[-1][1])) and parts \
            and not parts[-1][0].endswith(" ") and not text.startswith(" ")
        parts.append((" " + text if gap else text, spaced))
    return "".join(text for text, _ in parts)


def body_paragraphs(doc) -> list:
    """Every paragraph in document order, including those inside tables.

    ``doc.paragraphs`` yields only the top-level ``w:p`` children of the body, so
    any paragraph nested in a ``w:tbl`` is invisible to it. A document that lays
    an item out in a table therefore segments as if that content were absent -
    silently, because the item still exists and still carries whatever sits
    outside the table.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    tbl_tag, p_tag = qn("w:tbl"), qn("w:p")

    def walk(parent, element, seen):
        out = []
        for child in element.iterchildren():
            if child.tag == p_tag:
                out.append(Paragraph(child, parent))
            elif child.tag == tbl_tag:
                for row in Table(child, parent).rows:
                    for cell in row.cells:
                        if id(cell._tc) in seen:
                            continue
                        seen.add(id(cell._tc))
                        out.extend(walk(cell, cell._tc, seen))
        return out

    return walk(doc, doc.element.body, set())


def paragraph_images(paragraphs, doc) -> dict[int, list[bytes]]:
    """Image bytes per paragraph position, from inline and floating drawings."""
    from docx.oxml.ns import qn

    drawing_tags = (qn("wp:inline"), qn("wp:anchor"))
    blip_tag, embed_attr = qn("a:blip"), qn("r:embed")

    found = {}
    for i, para in enumerate(paragraphs):
        drawings = [d for tag in drawing_tags for d in para._element.findall(".//" + tag)]
        blips = [d.find(".//" + blip_tag) for d in drawings]
        ids = [b.get(embed_attr) for b in blips if b is not None]
        blobs = [doc.part.rels[r].target_part.blob
                 for r in ids if r and r in doc.part.rels]
        if blobs:
            found[i] = blobs
    return found


def segment_docx(
    source: str | Path | bytes,
    marker: str,
    label_position: str = "leading",
    skip_prefixes: Sequence[str] = (),
    section_markers: dict[str, str] | None = None,
) -> list[dict[str, Any]]:
    """Split a Word document into one record per item.

    Args:
        source: Path to the .docx file, or its raw bytes.
        marker: Line prefix carrying the item identifier, e.g. ``"Question ID:"``.
        label_position: ``"leading"`` when the marker introduces the block that
            follows it, ``"trailing"`` when it names the block above it.
        skip_prefixes: Line prefixes to drop, such as rule lines or headers.
        section_markers: Maps a header prefix to a section name. Text under such
            a header is stored in that section instead of the item body.

    Returns:
        One dict per identifier, holding ``item_id``, ``text``, ``images``, and
        any named sections, in document order.
    """
    if label_position not in label_positions:
        raise ValueError(f"label_position must be one of {label_positions}")
    if not docx_available:
        raise ImportError(
            "python-docx is required for .docx segmentation. "
            "Install with: pip install mathipy[documents]"
        )
    if isinstance(source, (bytes, bytearray)):
        doc = DocxDocument(io.BytesIO(bytes(source)))
    else:
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        doc = DocxDocument(str(path))
    paragraphs = body_paragraphs(doc)
    para_images = paragraph_images(paragraphs, doc)
    sections = section_markers or {}

    items: dict[str, dict[str, Any]] = {}
    order: list[str] = []
    buffer_text: list[str] = []
    buffer_images: list[bytes] = []
    buffer_section: str | None = None

    def record(item_id: str) -> dict[str, Any]:
        if item_id not in items:
            items[item_id] = {"item_id": item_id, "text": "", "images": [],
                              **{name: "" for name in sections.values()}}
            order.append(item_id)
        return items[item_id]

    def flush(item_id: str) -> None:
        nonlocal buffer_text, buffer_images, buffer_section
        item = record(item_id)
        if buffer_section:
            joined = " ".join(buffer_text)
            item[buffer_section] = f"{item[buffer_section]} {joined}".strip()
        else:
            item["text"] = " ".join(filter(None, [item["text"], *buffer_text]))
            item["images"].extend(buffer_images)
        buffer_text, buffer_images, buffer_section = [], [], None

    pending_id: str | None = None
    for i, para in enumerate(paragraphs):
        text = paragraph_text(para).strip()

        if text.startswith(marker):
            item_id = text[len(marker):].strip()
            if label_position == "trailing":
                flush(item_id)
            else:
                if pending_id is not None:
                    flush(pending_id)
                else:
                    buffer_text, buffer_images, buffer_section = [], [], None
                pending_id = item_id
                record(item_id)
            continue

        matched = next((name for prefix, name in sections.items()
                        if text.startswith(prefix)), None)
        if matched:
            buffer_section = matched
            continue

        if skip_prefixes and text.startswith(tuple(skip_prefixes)):
            continue

        if text:
            buffer_text.append(text)
        buffer_images.extend(para_images.get(i, []))

    if label_position == "leading" and pending_id is not None:
        flush(pending_id)

    return [items[k] for k in order]


def check_alignment(
    items: Sequence[dict[str, Any]],
    reference: dict[str, str],
    id_key: str = "item_id",
    text_key: str = "text",
    min_length: int = 25,
) -> dict[str, Any]:
    """Report whether each item's content matches its own reference text.

    Compares every item against the reference for itself, the item before, and
    the item after. A majority matching a neighbour means the identifiers are
    offset by one position.

    Returns:
        Counts and proportions for ``own``, ``next``, and ``previous``, plus a
        ``verdict`` of ``"aligned"``, ``"shifted_next"``, ``"shifted_previous"``,
        or ``"undetermined"``.
    """
    ids = [it.get(id_key) for it in items]
    tallies = {"own": 0, "next": 0, "previous": 0}
    compared = 0

    for n, item in enumerate(items):
        body = (item.get(text_key) or "").strip()
        if len(body) < min_length:
            continue
        scores = {
            "own": _overlap(body, reference.get(ids[n], "")),
            "next": _overlap(body, reference.get(ids[n + 1], "")) if n + 1 < len(ids) else 0.0,
            "previous": _overlap(body, reference.get(ids[n - 1], "")) if n else 0.0,
        }
        best = max(scores.values())
        if best == 0:
            continue
        compared += 1
        tallies[max(scores, key=scores.get)] += 1

    if not compared:
        return {**tallies, "compared": 0, "verdict": "undetermined"}

    shares = {k: v / compared for k, v in tallies.items()}
    leader = max(shares, key=shares.get)
    verdict = {
        "own": "aligned",
        "next": "shifted_next",
        "previous": "shifted_previous",
    }[leader] if shares[leader] > 0.5 else "undetermined"

    return {**tallies, "compared": compared,
            "proportions": {k: round(v, 4) for k, v in shares.items()},
            "verdict": verdict}
