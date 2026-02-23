"""Tests for mathipy.ocr multi-format extraction module."""

import importlib.util
import os
from pathlib import Path
from unittest.mock import patch

import pytest

from mathipy.ocr import (
    MultimodalOCR,
    _empty_result,
    _extract_answer_choices,
    _parse_response,
)
from mathipy.utils import extract_math_expressions

data_dir = Path(os.environ.get("MATHIPY_TEST_DATA", Path(__file__).resolve().parent / "data"))
images_dir = data_dir / "images"


def _has_api_key():
    from mathipy._api import _load_dotenv
    _load_dotenv()
    return bool(os.getenv("GEMINI_API_KEY") or os.getenv("OPENAI_API_KEY"))


needs_api = pytest.mark.skipif(not _has_api_key(), reason="No API key set")


@pytest.fixture
def sample_image_path():
    images = sorted(images_dir.glob("*.png"))
    if not images:
        pytest.skip("No PNG images found in data/images")
    return images[0]


@pytest.fixture
def sample_txt_path(tmp_path):
    content = (
        "What is the value of 3 + 4 × 2?\n"
        "\n"
        "A) 10\n"
        "B) 11\n"
        "C) 14\n"
        "D) 7\n"
        "\n"
        "The expression $3 + 4 \\times 2 = 11$.\n"
    )
    p = tmp_path / "sample.txt"
    p.write_text(content, encoding="utf-8")
    return p


@pytest.fixture
def sample_docx_path(tmp_path):
    try:
        from docx import Document as DocxDocument
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = DocxDocument()
    doc.add_paragraph("A rectangle has a length of 8 cm and a width of 5 cm.")
    doc.add_paragraph("What is the area of the rectangle?")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Length"
    table.cell(0, 1).text = "8 cm"
    table.cell(1, 0).text = "Width"
    table.cell(1, 1).text = "5 cm"

    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return p


@pytest.fixture
def sample_pdf_path(tmp_path):
    if not importlib.util.find_spec("pdfplumber"):
        pytest.skip("pdfplumber not installed")

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed (needed to create test PDF)")

    p = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 700, "Solve for x: 2x + 5 = 15")
    c.drawString(72, 680, "A) x = 5")
    c.drawString(72, 660, "B) x = 10")
    c.drawString(72, 640, "C) x = 7.5")
    c.save()
    return p


# --- Initialization ---

class TestInit:
    def test_lazy_init_no_api_key(self):
        ocr = MultimodalOCR()
        assert ocr.provider == "gemini"
        assert ocr.model == "gemini-2.5-flash"

    def test_lazy_init_no_httpx(self):
        with patch.dict("sys.modules", {"httpx": None}):
            ocr = MultimodalOCR()
            assert ocr.provider == "gemini"

    def test_provider_validation(self):
        with pytest.raises(ValueError, match="Unsupported provider"):
            MultimodalOCR(provider="claude")

    def test_mode_validation(self):
        ocr = MultimodalOCR()
        with pytest.raises(ValueError, match="Unsupported mode"):
            ocr.extract("dummy.txt", mode="summary")


# --- API Key Resolution ---

class TestAPIKey:
    def test_resolve_api_key_from_override(self):
        ocr = MultimodalOCR(api_key="test-key-123")
        assert ocr._resolve_api_key() == "test-key-123"

    def test_resolve_api_key_missing(self):
        ocr = MultimodalOCR()
        with patch.dict(os.environ, {}, clear=True):
            with pytest.raises(ValueError, match="API key not found"):
                ocr._resolve_api_key()


# --- Text Extraction ---

class TestExtractTxt:
    def test_extract_txt(self, sample_txt_path):
        ocr = MultimodalOCR()
        result = ocr.extract(sample_txt_path)

        assert result["extraction_confidence"] == 1.0
        assert result["content_type"] == "text_only"
        assert "3 + 4" in result["full_text"]
        assert len(result["answer_choices"]) > 0
        assert len(result["math_expressions"]) > 0

    def test_extract_txt_file_not_found(self):
        ocr = MultimodalOCR()
        with pytest.raises(FileNotFoundError):
            ocr.extract("nonexistent.txt")


# --- Routing ---

class TestRouting:
    def test_extract_routing_txt(self, sample_txt_path):
        ocr = MultimodalOCR()
        with patch.object(ocr, "_extract_from_txt", wraps=ocr._extract_from_txt) as mock:
            ocr.extract(sample_txt_path)
            mock.assert_called_once()

    def test_extract_routing_bytes(self):
        ocr = MultimodalOCR()
        with patch.object(ocr, "_extract_from_image") as mock:
            mock.return_value = _empty_result()
            ocr.extract(b"\x89PNG\r\n\x1a\n")
            mock.assert_called_once()
            args = mock.call_args[0]
            assert isinstance(args[0], bytes)

    def test_extract_routing_url(self):
        ocr = MultimodalOCR()
        with patch.object(ocr, "_extract_from_image") as mock:
            mock.return_value = _empty_result()
            ocr.extract("https://example.com/image.png")
            mock.assert_called_once()
            assert mock.call_args[0][0] == "https://example.com/image.png"

    def test_extract_routing_image_ext(self, sample_image_path):
        ocr = MultimodalOCR()
        with patch.object(ocr, "_extract_from_image") as mock:
            mock.return_value = _empty_result()
            ocr.extract(sample_image_path)
            mock.assert_called_once()

    def test_extract_routing_docx(self, tmp_path):
        ocr = MultimodalOCR()
        fake_docx = tmp_path / "test.docx"
        fake_docx.touch()
        with patch.object(ocr, "_extract_from_docx") as mock:
            mock.return_value = _empty_result()
            ocr.extract(fake_docx)
            mock.assert_called_once()

    def test_extract_routing_pdf(self, tmp_path):
        ocr = MultimodalOCR()
        fake_pdf = tmp_path / "test.pdf"
        fake_pdf.touch()
        with patch.object(ocr, "_extract_from_pdf") as mock:
            mock.return_value = _empty_result()
            ocr.extract(fake_pdf)
            mock.assert_called_once()


# --- Docx Extraction ---

class TestExtractDocx:
    def test_extract_docx_text_only(self, sample_docx_path):
        ocr = MultimodalOCR()
        result = ocr.extract(sample_docx_path)

        assert result["content_type"] == "text_only"
        assert result["extraction_confidence"] == 0.95
        assert "rectangle" in result["full_text"].lower()
        assert 8.0 in result["numbers_found"]

    def test_extract_docx_with_tables(self, sample_docx_path):
        ocr = MultimodalOCR()
        result = ocr.extract(sample_docx_path)

        assert len(result["data_elements"]) > 0
        table_text = " ".join(result["data_elements"])
        assert "Length" in table_text
        assert "Width" in table_text

    def test_extract_docx_missing_dep(self):
        ocr = MultimodalOCR()
        with patch("mathipy.ocr.docx_available", False):
            with pytest.raises(ImportError, match="python-docx"):
                ocr._extract_from_docx(Path("test.docx"), "full", 30)


# --- PDF Extraction ---

class TestExtractPdf:
    def test_extract_pdf_text_only(self, sample_pdf_path):
        ocr = MultimodalOCR()
        result = ocr.extract(sample_pdf_path)

        assert result["content_type"] == "text_only"
        assert result["extraction_confidence"] == 0.95
        assert "x" in result["full_text"].lower()

    def test_extract_pdf_missing_dep(self):
        ocr = MultimodalOCR()
        with patch("mathipy.ocr.pdfplumber_available", False):
            with pytest.raises(ImportError, match="pdfplumber"):
                ocr._extract_from_pdf(Path("test.pdf"), "full", 30)


# --- Merge Image Results ---

class TestMerge:
    def test_merge_image_results(self):
        ocr = MultimodalOCR()
        base = _empty_result()
        base["full_text"] = "Base text"
        base["extraction_confidence"] = 0.95
        base["math_expressions"] = ["x = 5"]
        base["numbers_found"] = ["5"]

        img_results = [
            {
                "image_description": "A triangle",
                "math_expressions": ["x = 5", "y = 10"],
                "labels": ["vertex A"],
                "numbers_found": ["5", "10"],
                "variables_found": ["x", "y"],
                "data_elements": ["row1"],
                "answer_choices": {"A": "5"},
                "extraction_confidence": 0.8,
            },
            {
                "image_description": "A circle",
                "math_expressions": ["r = 3"],
                "labels": ["center"],
                "numbers_found": ["3"],
                "variables_found": ["r"],
                "data_elements": ["row2"],
                "answer_choices": {"B": "10"},
                "extraction_confidence": 0.9,
            },
        ]

        result = ocr._merge_image_results(base, img_results)

        assert result["content_type"] == "mixed"
        assert "[Image 1] A triangle" in result["image_description"]
        assert "[Image 2] A circle" in result["image_description"]

        assert "x = 5" in result["math_expressions"]
        assert "y = 10" in result["math_expressions"]
        assert "r = 3" in result["math_expressions"]

        assert "5" in result["numbers_found"]
        assert "10" in result["numbers_found"]
        assert "3" in result["numbers_found"]

        assert result["answer_choices"]["A"] == "5"
        assert result["answer_choices"]["B"] == "10"

        expected_conf = round(0.7 * 0.95 + 0.3 * 0.85, 2)
        assert result["extraction_confidence"] == expected_conf

    def test_merge_preserves_base_text(self):
        ocr = MultimodalOCR()
        base = _empty_result()
        base["full_text"] = "Original text preserved"
        base["extraction_confidence"] = 0.95

        img_results = [{"image_description": "diagram", "extraction_confidence": 0.7}]
        result = ocr._merge_image_results(base, img_results)

        assert result["full_text"] == "Original text preserved"


# --- Helper Functions ---

class TestHelpers:
    def test_empty_result_template(self):
        result = _empty_result()
        expected_keys = {
            "content_type", "full_text", "image_description",
            "question_text", "math_expressions", "answer_choices",
            "data_elements", "labels", "numbers_found",
            "variables_found", "extraction_confidence",
        }
        assert set(result.keys()) == expected_keys
        assert result["content_type"] == "text_only"
        assert result["extraction_confidence"] == 0.0
        assert result["math_expressions"] == []
        assert result["answer_choices"] == {}

    def test_extract_math_expressions(self):
        text = "The equation $x^2 + 1 = 0$ and \\frac{1}{2} is given."
        exprs = extract_math_expressions(text)
        assert any("x^2" in e for e in exprs)

    def test_extract_answer_choices(self):
        text = "A) 10  B) 20  C) 30  D) 40"
        choices = _extract_answer_choices(text)
        assert "A" in choices
        assert "10" in choices["A"]

    def test_parse_response_json(self):
        raw = '```json\n{"content_type": "text_only", "full_text": "hello"}\n```'
        result = _parse_response(raw)
        assert result["content_type"] == "text_only"
        assert result["full_text"] == "hello"

    def test_parse_response_plain_text(self):
        result = _parse_response("This is plain text with numbers 42")
        assert result["content_type"] == "text_only"
        assert 42.0 in result["numbers_found"]


# --- Convenience Functions ---

class TestClassBasedAPI:
    def test_multimodal_ocr_extract_method(self, sample_txt_path):
        result = MultimodalOCR().extract(sample_txt_path)
        assert result["extraction_confidence"] == 1.0
        assert "3 + 4" in result["full_text"]

    def test_extract_returns_dict(self, sample_txt_path):
        result = MultimodalOCR().extract(sample_txt_path)
        assert result["extraction_confidence"] == 1.0
        assert result["content_type"] == "text_only"


# --- API-Dependent Tests ---

class TestImageAPI:
    @needs_api
    def test_image_full_extraction(self, sample_image_path):
        ocr = MultimodalOCR()
        result = ocr.extract(sample_image_path, mode="full")

        assert result["content_type"] in ("text_only", "image_only", "mixed")
        has_content = (
            bool((result.get("full_text") or "").strip())
            or bool((result.get("image_description") or "").strip())
        )
        assert has_content, "Expected non-empty full_text or image_description"
        assert 0 < result["extraction_confidence"] <= 1.0

    @needs_api
    def test_image_describe_mode(self, sample_image_path):
        ocr = MultimodalOCR()
        result = ocr.extract(sample_image_path, mode="describe")

        has_content = (
            bool((result.get("full_text") or "").strip())
            or bool((result.get("image_description") or "").strip())
        )
        assert has_content, "Describe mode should return content"

    @needs_api
    def test_extract_with_image(self, sample_image_path):
        result = MultimodalOCR().extract(str(sample_image_path))
        assert isinstance(result, dict)
        assert "full_text" in result
        assert "extraction_confidence" in result
